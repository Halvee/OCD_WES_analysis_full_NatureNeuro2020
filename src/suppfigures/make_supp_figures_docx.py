from docx import Document
from docx.shared import Inches
import argparse

def main():

    # parse user-defined args
    args = parse_args()

    # get headers and image files from userinput
    figurenames = []
    imgfiles = []
    figinfo = {}
    for i in range(len(args.figurenames_imgfiles)):
        figurename_imgfile = args.figurenames_imgfiles[i].split("%")
        figurename=figurename_imgfile[0]
        width_inches=float(figurename_imgfile[1])
        imgfile=figurename_imgfile[2]
        figurenames.append(figurename)
        imgfiles.append(imgfile)
        if figurename not in figinfo: 
            figinfo[figurename]={"imgfiles":[],
                                 "widths_inches":[],
                                 "desc":"",
                                 "paragraph":""}
        figinfo[figurename]["imgfiles"].append(imgfile)
        figinfo[figurename]["widths_inches"].append(width_inches)
    
    # load captions if any are defined by user
    figinfo = read_captions_txt(args.captions_txt, figinfo)

    # init Document object for creating the docx file
    document = Document()

    # create header on first page, but leave table of contents blank. Works best
    # to add manually afterwards, as MS word will be able to link header on each
    # page along with page number
    document.add_heading(args.header, 0)
    document.add_page_break()

    # create table of contents on first page
    #document.add_heading('Supplemental Figures', 0)
    #prev_figurename=None
    #for i in range(len(figurenames)):
    #    figurename = figurenames[i]
    #    if figurename != prev_figurename:
    #        header = figurename
    #        if figinfo[figurename]["desc"] != "":
    #            header = figurename + " : " + figinfo[figurename]["desc"]
    #        document.add_heading(header,level=1)
    #    prev_figurename = figurename
    #document.add_page_break()

    # add imgfiles to captions
    prev_figurenames=set([])
    for i in range(len(figurenames)):
        figurename = figurenames[i]
        if figurename in prev_figurenames:
            continue
        header = figurename
        if figinfo[figurename]["desc"] != "":
            header = figurename + " : " + figinfo[figurename]["desc"]
        document.add_heading(header,level=1)
        for j in range(len(figinfo[figurename]["imgfiles"])):
            imgfile = figinfo[figurename]["imgfiles"][j]
            width_inches = figinfo[figurename]["widths_inches"][j]
            document.add_picture(imgfile,width=Inches(width_inches))
        if figinfo[figurename]["paragraph"] != "":
            document.add_paragraph(figinfo[figurename]["paragraph"])
        if i < (len(figurenames)-1): document.add_page_break()
        prev_figurenames.add(figurename)
        
    # save the document
    document.save(args.out_docx)

    return

def read_captions_txt(captions_txt_file, figinfo):
    if captions_txt_file == None:
        return figinfo
    fh = open(captions_txt_file,"r")
    current_figure_name=None
    for line in fh:
        line = line.rstrip()
        if line[:2] == "##":
            info = line[2:].split(":")
            figure_name = info[0]
            figure_desc = info[1]
            if figure_name in figinfo:
                figinfo[figure_name]["desc"] = figure_desc
            current_figure_name=figure_name
        else:
            line=line+" "
            figinfo[current_figure_name]["paragraph"]+=line
    fh.close()
    return figinfo

def parse_args():
    parser = argparse.ArgumentParser(description='Make the supplemental figures docx file.')
    parser.add_argument('--out-docx', action='store',
                        type=str, default='out.docx',
                        help='outpout docx filename.')
    parser.add_argument('--header',action='store',type=str,
                        default='Supplemental Figures',
                        help='text header for top of supplemental doc')
    parser.add_argument('--captions-txt', action='store', type=str, 
                        default=None, 
                        help='text file that stores captions to insert ' + \
                             'each input supplemental figure')
    parser.add_argument('--width-inches', action='store', type=str,
                        default=None,
                        help='with of each individual figure inserted ' + \
                             'into file. Format : inchesX_numX,inchesY_numY')
    parser.add_argument('figurenames_imgfiles', type=str, nargs='+',
                        help='header and imgfile, seperated by a %.')
    args = parser.parse_args()
    return args

if __name__ == "__main__":
    main()
